# Projeto 5 - Detecção Automática de Anomalias em Logs de Pipelines de Dados
# Módulo de Análise de Arquivo

# Imports
import pandas as pd
from docx import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_community.llms.ollama import Ollama

# Instanciação do LLM Llama3 através do Ollama
llm = Ollama(model="llama3")

# Criação do parser para a saída do modelo de linguagem
output_parser = StrOutputParser()

# Função para gerar texto baseado nos dados do Excel
def dsa_gera_resultados():

    # Carrega os dados do Excel
    df = pd.read_excel('logs_pipeline.xlsx')

    # Inicializa uma lista para armazenar os resultados
    resultados = []

    # Criação do template de prompt 
    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", "Você é um analista de dados especializado. Analise os dados de execução do pipeline e forneça feedback em português do Brasil."),
            ("user", "question: {question}")
        ]
    )

    # Definição da cadeia de execução: prompt -> LLM -> output_parser
    chain = prompt | llm | output_parser

    # Criação do documento DOCX
    document = Document()
    document.add_heading('Resultados da Análise de Logs do Pipeline', 0)

    # Itera sobre as linhas do DataFrame
    for _, row in df.iterrows():
        
        # Extrai os valores de cada linha
        pipeline_id, execution_id, start_time, end_time, status, execution_time_minutes, operating_system, processing_type, attempt_number = row
        
        # Cria o prompt para o LLM com base nos dados da execução do pipeline
        consulta_pipeline = (f"Pipeline ID: {pipeline_id}, Execution ID: {execution_id}, Status: {status}, Tempo de Execução: {execution_time_minutes} minutos,"
                             f"Operating_System: {operating_system}, Processing_Type: {processing_type}, Attempt_Number: {attempt_number}")
        
        # Gera o texto de resultado usando o LLM
        response = chain.invoke({'question': consulta_pipeline})
        
        # Adiciona o texto gerado à lista de resultados
        resultados.append(response)
        
        # Adiciona o resultado ao documento DOCX
        document.add_paragraph(response)

    # Salva os resultados em um arquivo DOCX
    document.save('projeto5-resultado.docx')

    # Retorna a lista de resultados
    return resultados

# Gera resultados chamando a função definida
resultados = dsa_gera_resultados()

# Imprime cada resultado gerado
for resultado in resultados:
    print(resultado)
